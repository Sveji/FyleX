import io
import os
import requests
import docx
import pymupdf
import striprtf.striprtf as striprtf
from urllib.parse import urlparse
from pathlib import Path


def extract_text_from_url(url):
    """
    Extract text from documents (.doc, .docx, .rtf, .txt, .pdf) stored at a cloud URL,
    preserving the original formatting as much as possible.

    Args:
        url (str): The URL to the document file.
                   Example: https://res.cloudinary.com/djm6yhqvx/raw/upload/v1742495357/hackTues11/filename.type

    Returns:
        str: The extracted text from the document with formatting preserved.

    Raises:
        ValueError: If the file type is not supported or if the URL doesn't contain a valid file extension.
        requests.RequestException: If there's an issue fetching the file from the URL.
    """

    parsed_url = urlparse(url)
    path = parsed_url.path
    filename = os.path.basename(path)
    file_extension = os.path.splitext(filename)[1].lower()


    response = requests.get(url)
    response.raise_for_status()
    file_content = response.content


    file_bytes = io.BytesIO(file_content)

    if file_extension == '.txt':
        return extract_text_from_txt(file_bytes)
    elif file_extension == '.pdf':
        return extract_text_from_pdf(file_bytes)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_bytes)
    elif file_extension == '.doc':
        return extract_text_from_doc(file_bytes)
    elif file_extension == '.rtf':
        return extract_text_from_rtf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def extract_text_from_txt(file_bytes):
    """
    Extract text from a .txt file, preserving line breaks.

    Args:
        file_bytes (BytesIO): The file content as BytesIO object.

    Returns:
        str: The extracted text with original formatting.
    """
    file_bytes.seek(0)
    content = file_bytes.read().decode('utf-8', errors='replace')

    return content


def extract_text_from_pdf(file_bytes):
    """
    Extract text from a PDF file using PyMuPDF with enhanced formatting preservation.

    Args:
        file_bytes (BytesIO): The file content as BytesIO object.

    Returns:
        str: The extracted text with formatting preserved.
    """
    file_bytes.seek(0)
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    full_text = []

    for page_num in range(doc.page_count):
        page = doc[page_num]
        blocks = page.get_text("dict")["blocks"]
        page_text = []

        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    line_text = []
                    for span in line["spans"]:
                        line_text.append(span["text"])
                    if line_text:
                        page_text.append(" ".join(line_text))
                if page_text:
                    page_text.append("")

        full_text.append("\n".join(page_text))

    doc.close()
    return "\n\n".join(full_text)


def extract_text_from_docx(file_bytes):
    """
    Extract text from a .docx file with enhanced formatting preservation.

    Args:
        file_bytes (BytesIO): The file content as BytesIO object.

    Returns:
        str: The extracted text with formatting preserved.
    """
    file_bytes.seek(0)
    doc = docx.Document(file_bytes)
    result = []


    for para in doc.paragraphs:
        if para.text.strip():
            if para.style.name.startswith('Heading'):
                result.append(para.text)
                result.append("")
            else:
                result.append(para.text)
        else:
            result.append("")

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(" | ".join(row_data))

        result.append("")
        result.extend(table_data)
        result.append("")

    return "\n".join(result)


def extract_text_from_doc(file_bytes):
    """
    Extract text from a .doc file using antiword (if available) or fallback to PDF conversion,
    with improved formatting preservation.

    Args:
        file_bytes (BytesIO): The file content as BytesIO object.

    Returns:
        str: The extracted text with formatting preserved where possible.
    """
    try:
        import subprocess
        import tempfile


        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
            temp_filename = temp_file.name
            file_bytes.seek(0)
            temp_file.write(file_bytes.read())

        try:

            result = subprocess.run(['antiword', '-f', temp_filename],  # -f preserves some formatting
                                    capture_output=True,
                                    text=True,
                                    check=True)
            extracted_text = result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            try:
                result = subprocess.run(['soffice', '--headless', '--convert-to', 'pdf',
                                         '--outdir', os.path.dirname(temp_filename),
                                         temp_filename],
                                        capture_output=True,
                                        check=True)

                pdf_path = os.path.splitext(temp_filename)[0] + '.pdf'

                with open(pdf_path, 'rb') as pdf_file:
                    pdf_bytes = io.BytesIO(pdf_file.read())
                    extracted_text = extract_text_from_pdf(pdf_bytes)

                os.remove(pdf_path)
            except:
                try:
                    result = subprocess.run(['catdoc', temp_filename],
                                            capture_output=True,
                                            text=True,
                                            check=True)
                    extracted_text = result.stdout
                except:
                    extracted_text = "Could not extract text from .doc file. Required tools not available."
        finally:
            os.remove(temp_filename)

        return extracted_text
    except Exception as e:
        return f"Error extracting text from .doc file: {str(e)}"


def extract_text_from_rtf(file_bytes):
    """
    Extract text from an RTF file with improved formatting preservation.

    Args:
        file_bytes (BytesIO): The file content as BytesIO object.

    Returns:
        str: The extracted text with formatting preserved where possible.
    """
    file_bytes.seek(0)
    rtf_content = file_bytes.read().decode('utf-8', errors='replace')
    plain_text = striprtf.rtf_to_text(rtf_content)

    import re
    formatted_text = re.sub(r'\n{3,}', '\n\n', plain_text)

    return formatted_text



url = "https://res.cloudinary.com/djm6yhqvx/raw/upload/v1742495357/hackTues11/iug6djyavyxphxnhwgdr.pdf"

extracted_text = extract_text_from_url(url)
print(extracted_text)